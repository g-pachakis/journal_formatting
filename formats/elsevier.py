"""
Elsevier journal format plugin.

Builds a formatted .docx from scratch using python-docx.
Target: single-column Elsevier submission format.
"""

import re
from docx import Document as DocxDocument
from docx.oxml.ns import qn as docx_qn
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree as lxml_etree
from caption_fields import add_caption_with_seq

FORMAT_NAME = 'Elsevier'
FORMAT_SUFFIX = '_Elsevier'

FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(12)
LINE_SPACING = 1.5
PAGE_MARGIN = Cm(2.5)

TITLE_SIZE = Pt(16)
AUTHOR_SIZE = Pt(12)
AFFILIATION_SIZE = Pt(10)
HEADING1_SIZE = Pt(13)
HEADING2_SIZE = Pt(12)
HEADING3_SIZE = Pt(12)
ABSTRACT_LABEL_SIZE = Pt(13)
KEYWORDS_SIZE = Pt(11)
TABLE_FONT_SIZE = Pt(10)
TABLE_CAPTION_SIZE = Pt(10)
REF_SIZE = Pt(11)


def _set_run_font(run, size=None, bold=None, italic=None, color=None,
                  superscript=False, subscript=False):
    run.font.name = FONT_NAME
    if size:
        run.font.size = size
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color
    if superscript:
        run.font.superscript = True
    if subscript:
        run.font.subscript = True


def _set_spacing(para, before=0, after=0, line_spacing=LINE_SPACING):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def _add_styled_para(doc, text, size=FONT_SIZE, bold=False, italic=False,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0,
                     first_line_indent=None, color=None):
    para = doc.add_paragraph()
    para.alignment = alignment
    _set_spacing(para, before, after)
    if first_line_indent:
        para.paragraph_format.first_line_indent = first_line_indent
    run = para.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return para


def _add_runs_para(doc, runs_data, size=FONT_SIZE,
                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   before=0, after=0, first_line_indent=None,
                   bold_override=None):
    para = doc.add_paragraph()
    para.alignment = alignment
    _set_spacing(para, before, after)
    if first_line_indent:
        para.paragraph_format.first_line_indent = first_line_indent
    for rd in runs_data:
        if not rd['text']:
            continue
        run = para.add_run(rd['text'])
        b = bold_override if bold_override is not None else rd['bold']
        _set_run_font(run, size=size, bold=b, italic=rd['italic'],
                      superscript=rd.get('superscript', False),
                      subscript=rd.get('subscript', False))
    return para


def _add_three_line_table(doc, table_data):
    if not table_data or not table_data[0]:
        return None
    ncols = sum(c.get('gridspan', 1) for c in table_data[0])
    nrows = len(table_data)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl_pr = table._tbl.find(docx_qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = lxml_etree.SubElement(table._tbl, docx_qn('w:tblPr'))

    borders = lxml_etree.SubElement(tbl_pr, docx_qn('w:tblBorders'))
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = lxml_etree.SubElement(borders, docx_qn(f'w:{border_name}'))
        el.set(docx_qn('w:val'), 'none')
        el.set(docx_qn('w:sz'), '0')
        el.set(docx_qn('w:space'), '0')
        el.set(docx_qn('w:color'), 'auto')

    top_border = borders.find(docx_qn('w:top'))
    top_border.set(docx_qn('w:val'), 'single')
    top_border.set(docx_qn('w:sz'), '12')
    top_border.set(docx_qn('w:color'), '000000')

    bottom_border = borders.find(docx_qn('w:bottom'))
    bottom_border.set(docx_qn('w:val'), 'single')
    bottom_border.set(docx_qn('w:sz'), '12')
    bottom_border.set(docx_qn('w:color'), '000000')

    for ri, row_data in enumerate(table_data):
        col_idx = 0
        for cell_data in row_data:
            if col_idx >= ncols:
                break
            if cell_data.get('vmerge_continue'):
                col_idx += cell_data.get('gridspan', 1)
                continue
            cell = table.cell(ri, col_idx)
            for p in cell.paragraphs:
                p.clear()
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if cell_data.get('runs'):
                for rd in cell_data['runs']:
                    if not rd['text']:
                        continue
                    run = para.add_run(rd['text'])
                    b = rd['bold'] or (ri == 0)
                    _set_run_font(run, size=TABLE_FONT_SIZE, bold=b,
                                  italic=rd['italic'],
                                  superscript=rd.get('superscript', False),
                                  subscript=rd.get('subscript', False))
            else:
                run = para.add_run(cell_data.get('text', '').strip())
                _set_run_font(run, size=TABLE_FONT_SIZE, bold=(ri == 0))
            gs = cell_data.get('gridspan', 1)
            if gs > 1 and col_idx + gs <= ncols:
                cell.merge(table.cell(ri, col_idx + gs - 1))
            tc_pr = cell._tc.find(docx_qn('w:tcPr'))
            if tc_pr is None:
                tc_pr = lxml_etree.SubElement(cell._tc, docx_qn('w:tcPr'))
            margins = lxml_etree.SubElement(tc_pr, docx_qn('w:tcMar'))
            for side, val in [('top', '40'), ('bottom', '40'),
                              ('left', '80'), ('right', '80')]:
                m = lxml_etree.SubElement(margins, docx_qn(f'w:{side}'))
                m.set(docx_qn('w:w'), val)
                m.set(docx_qn('w:type'), 'dxa')
            col_idx += gs

    if nrows > 1:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.find(docx_qn('w:tcPr'))
            if tc_pr is None:
                tc_pr = lxml_etree.SubElement(cell._tc, docx_qn('w:tcPr'))
            tc_borders = lxml_etree.SubElement(tc_pr, docx_qn('w:tcBorders'))
            bottom = lxml_etree.SubElement(tc_borders, docx_qn('w:bottom'))
            bottom.set(docx_qn('w:val'), 'single')
            bottom.set(docx_qn('w:sz'), '8')
            bottom.set(docx_qn('w:space'), '0')
            bottom.set(docx_qn('w:color'), '000000')

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                pf = para.paragraph_format
                pf.space_before = Pt(1)
                pf.space_after = Pt(1)
                pf.line_spacing = 1.15
    return table


def build(items, output_path, ris_data=None, zotero_enabled=False):
    doc = DocxDocument()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = PAGE_MARGIN
    section.right_margin = PAGE_MARGIN
    section.top_margin = PAGE_MARGIN
    section.bottom_margin = PAGE_MARGIN

    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.line_spacing = LINE_SPACING
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    abstract_text_parts = []
    keywords_text = ''
    for elem in items:
        if elem['type'] == 'abstract_text':
            abstract_text_parts.append(elem)
        if elem['type'] == 'keywords':
            keywords_text = elem['text']

    title_elem = None
    for elem in items:
        if elem['type'] in ('abstract_heading', 'heading1'):
            break
        if elem['type'] == 'paragraph':
            if not title_elem:
                title_elem = elem

    title_text = title_elem['text'] if title_elem else '[Title]'
    _add_styled_para(doc, title_text, size=TITLE_SIZE, bold=True,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, after=12)

    _add_styled_para(doc, '', after=6)
    _add_styled_para(doc, '[Author Name]a,*, [Author Name]b',
                     size=AUTHOR_SIZE,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    _add_styled_para(doc,
                     'a [Department, University/Institution, City, Country]',
                     size=AFFILIATION_SIZE, italic=True,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    _add_styled_para(doc,
                     'b [Department, University/Institution, City, Country]',
                     size=AFFILIATION_SIZE, italic=True,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    _add_styled_para(
        doc,
        '* Corresponding author. E-mail address: [email@institution.edu]',
        size=AFFILIATION_SIZE, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, after=18)

    _add_styled_para(doc, 'Highlights', size=HEADING1_SIZE, bold=True,
                     before=12, after=6)
    for h in ['[First key finding — max 85 characters per bullet]',
              '[Second key finding]', '[Third key finding]',
              '[Fourth key finding (optional)]',
              '[Fifth key finding (optional)]']:
        para = doc.add_paragraph()
        para.style = doc.styles['List Bullet']
        run = para.add_run(h)
        _set_run_font(run, size=FONT_SIZE)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = LINE_SPACING

    _add_styled_para(doc, '', after=6)
    _add_styled_para(doc, 'Graphical Abstract', size=HEADING1_SIZE,
                     bold=True, before=12, after=6)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(para, before=6, after=6)
    run = para.add_run(
        '[Insert graphical abstract here — recommended size: '
        '531 x 1328 pixels (min. 300 DPI)]')
    _set_run_font(run, size=Pt(10), italic=True,
                  color=RGBColor(0x80, 0x80, 0x80))

    ga_table = doc.add_table(rows=1, cols=1)
    ga_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = ga_table.cell(0, 0)
    tc_pr = cell._tc.find(docx_qn('w:tcPr'))
    if tc_pr is None:
        tc_pr = lxml_etree.SubElement(cell._tc, docx_qn('w:tcPr'))
    tc_w = lxml_etree.SubElement(tc_pr, docx_qn('w:tcW'))
    tc_w.set(docx_qn('w:w'), '9072')
    tc_w.set(docx_qn('w:type'), 'dxa')
    tc_borders = lxml_etree.SubElement(tc_pr, docx_qn('w:tcBorders'))
    for side in ['top', 'bottom', 'left', 'right']:
        b = lxml_etree.SubElement(tc_borders, docx_qn(f'w:{side}'))
        b.set(docx_qn('w:val'), 'dashed')
        b.set(docx_qn('w:sz'), '6')
        b.set(docx_qn('w:space'), '0')
        b.set(docx_qn('w:color'), '999999')
    tc_mar = lxml_etree.SubElement(tc_pr, docx_qn('w:tcMar'))
    for side, val in [('top', '1200'), ('bottom', '1200'),
                      ('left', '120'), ('right', '120')]:
        m = lxml_etree.SubElement(tc_mar, docx_qn(f'w:{side}'))
        m.set(docx_qn('w:w'), val)
        m.set(docx_qn('w:type'), 'dxa')
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_run = cell_para.add_run('GRAPHICAL ABSTRACT')
    _set_run_font(cell_run, size=Pt(14), color=RGBColor(0xBB, 0xBB, 0xBB))

    _add_styled_para(doc, '', after=6)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(para, before=12, after=6)
    run = para.add_run('Abstract')
    _set_run_font(run, size=ABSTRACT_LABEL_SIZE, bold=True)

    for abs_elem in abstract_text_parts:
        _add_runs_para(doc, abs_elem['runs'], size=FONT_SIZE,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       before=0, after=6)

    if keywords_text:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_spacing(para, before=6, after=12)
        kw_label = para.add_run('Keywords: ')
        _set_run_font(kw_label, size=KEYWORDS_SIZE, italic=True)
        kw_text = keywords_text.replace('Keywords:', '').replace(
            'Keywords :', '').strip()
        kw_run = para.add_run(kw_text)
        _set_run_font(kw_run, size=KEYWORDS_SIZE, italic=True)

    skip_types = {'abstract_heading', 'abstract_text', 'keywords'}
    is_first_para = True
    found_first_heading = False

    for elem in items:
        etype = elem['type']
        if elem is title_elem:
            continue
        if etype in skip_types:
            continue

        if etype == 'heading1':
            found_first_heading = True
            is_first_para = True
            _add_styled_para(doc, elem['text'], size=HEADING1_SIZE,
                             bold=True, before=18, after=6)
            continue
        if etype == 'heading2':
            is_first_para = True
            _add_styled_para(doc, elem['text'], size=HEADING2_SIZE,
                             bold=True, italic=True, before=12, after=6)
            continue
        if etype == 'heading3':
            is_first_para = True
            _add_styled_para(doc, elem['text'], size=HEADING3_SIZE,
                             italic=True, before=10, after=4)
            continue
        if etype == 'references_heading':
            is_first_para = True
            _add_styled_para(doc, 'References', size=HEADING1_SIZE,
                             bold=True, before=18, after=6)
            continue
        if not found_first_heading:
            continue

        if etype == 'paragraph':
            indent = None if is_first_para else Cm(0.75)
            _add_runs_para(doc, elem['runs'], size=FONT_SIZE,
                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           before=0, after=6, first_line_indent=indent)
            is_first_para = False
            continue
        if etype == 'equation':
            _add_runs_para(doc, elem['runs'], size=FONT_SIZE,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           before=6, after=6)
            is_first_para = False
            continue
        if etype == 'table_caption':
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_spacing(para, before=12, after=4)
            text = elem['text']
            match = re.match(r'^Table\s+(\d+)\.?\s*(.*)', text)
            if match:
                num = match.group(1)
                desc = match.group(2)
                add_caption_with_seq(para, 'Table', num, description=desc,
                                     font_name=FONT_NAME, font_size=TABLE_CAPTION_SIZE,
                                     bold_label=True)
            else:
                run = para.add_run(text)
                _set_run_font(run, size=TABLE_CAPTION_SIZE)
            is_first_para = True
            continue
        if etype == 'table':
            _add_three_line_table(doc, elem['rows'])
            is_first_para = True
            continue
        if etype == 'table_footer':
            _add_styled_para(doc, elem['text'], size=Pt(9), italic=True,
                             before=2, after=6)
            is_first_para = True
            continue
        if etype == 'figure_placeholder':
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_spacing(para, before=12, after=12)
            text = elem['text']
            match = re.match(r'^\[?\s*Figure\s+(\d+)\]?\.?\s*(.*)', text, re.IGNORECASE)
            if match:
                num = match.group(1)
                desc = match.group(2)
                add_caption_with_seq(para, 'Figure', num, description=desc,
                                     font_name=FONT_NAME, font_size=TABLE_CAPTION_SIZE,
                                     bold_label=True)
            else:
                run = para.add_run(text)
                _set_run_font(run, size=TABLE_CAPTION_SIZE)
            is_first_para = True
            continue
        if etype == 'reference':
            p = _add_runs_para(doc, elem['runs'], size=REF_SIZE,
                               alignment=WD_ALIGN_PARAGRAPH.LEFT,
                               before=0, after=3,
                               first_line_indent=Cm(-0.75))
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.line_spacing = 1.15
            continue

    doc.save(output_path)
    return output_path
