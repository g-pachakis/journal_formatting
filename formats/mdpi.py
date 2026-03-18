"""
MDPI journal format plugin.

Builds a formatted .docx from scratch using python-docx.
All MDPI styles are registered as named Word styles.
"""

import re
from docx import Document as DocxDocument
from docx.oxml.ns import qn as docx_qn
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from lxml import etree as lxml_etree

from caption_fields import (register_mdpi_styles, add_caption_with_seq,
                             add_zotero_citation_field, ris_to_csl_json)
from reference_engine import resolve_references

FORMAT_NAME = 'MDPI'
FORMAT_SUFFIX = '_MDPI'

FONT_NAME = 'Palatino Linotype'
FONT_COLOR = RGBColor(0x00, 0x00, 0x00)
DEFAULT_SIZE = Pt(10)
LEFT_INDENT = Pt(130.4)

TABLE_WIDTH = 7857
TABLE_INDENT = 2608


def _set_run_font(run, size=None, bold=None, italic=None,
                  superscript=False, subscript=False):
    run.font.name = FONT_NAME
    run.font.color.rgb = FONT_COLOR
    if size:
        run.font.size = size
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if superscript:
        run.font.superscript = True
    if subscript:
        run.font.subscript = True


def _add_styled_para(doc, style_name, text='', bold_prefix=None,
                     font_size=None):
    """Add a paragraph with a named Word style."""
    para = doc.add_paragraph()
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass
    fs = font_size or DEFAULT_SIZE
    if bold_prefix:
        run = para.add_run(bold_prefix)
        _set_run_font(run, size=fs, bold=True)
    if text:
        run = para.add_run(text)
        _set_run_font(run, size=fs)
    return para


def _add_runs_para(doc, style_name, runs_data, bold_prefix=None,
                   font_size=None):
    """Add a paragraph with a named style, preserving inline formatting."""
    para = doc.add_paragraph()
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass
    fs = font_size or DEFAULT_SIZE
    if bold_prefix:
        run = para.add_run(bold_prefix)
        _set_run_font(run, size=fs, bold=True)
    for rd in runs_data:
        if not rd['text']:
            continue
        run = para.add_run(rd['text'])
        _set_run_font(run, size=fs, bold=rd.get('bold'), italic=rd.get('italic'),
                      superscript=rd.get('superscript', False),
                      subscript=rd.get('subscript', False))
    return para


def _apply_direct_style(para, font_size=None, bold=False, italic=False,
                        alignment=None, space_before=None, space_after=None,
                        line_spacing_pt=None, left_indent=None,
                        first_indent=None, outline_level=None):
    """Apply direct paragraph formatting (for styles not in doc.styles)."""
    if alignment is not None:
        para.alignment = alignment
    pf = para.paragraph_format
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing_pt is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        pf.line_spacing = line_spacing_pt
    if left_indent is not None:
        pf.left_indent = left_indent
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if outline_level is not None:
        ppr = para._p.find(docx_qn('w:pPr'))
        if ppr is None:
            ppr = lxml_etree.SubElement(para._p, docx_qn('w:pPr'))
        ol = lxml_etree.SubElement(ppr, docx_qn('w:outlineLvl'))
        ol.set(docx_qn('w:val'), str(outline_level))


def _build_author_placeholder(doc):
    para = doc.add_paragraph()
    _apply_direct_style(para, space_after=Pt(18),
                        line_spacing_pt=Pt(13))
    parts = [
        ('Firstname Lastname ', {'bold': True}),
        ('1', {'superscript': True}),
        (', Firstname Lastname ', {'bold': True}),
        ('2', {'superscript': True}),
        (' and Firstname Lastname ', {'bold': True}),
        ('2,', {'superscript': True}),
        ('*', {'bold': True}),
    ]
    for text, fmt in parts:
        run = para.add_run(text)
        _set_run_font(run, size=DEFAULT_SIZE, **fmt)
    return para


def _add_three_line_table(doc, table_data):
    if not table_data or not table_data[0]:
        return None
    ncols = sum(c.get('gridspan', 1) for c in table_data[0])
    if ncols == 0:
        return None
    nrows = len(table_data)
    table = doc.add_table(rows=nrows, cols=ncols)

    tbl_pr = table._tbl.find(docx_qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = lxml_etree.SubElement(table._tbl, docx_qn('w:tblPr'))

    tw = lxml_etree.SubElement(tbl_pr, docx_qn('w:tblW'))
    tw.set(docx_qn('w:w'), str(TABLE_WIDTH))
    tw.set(docx_qn('w:type'), 'dxa')
    ti = lxml_etree.SubElement(tbl_pr, docx_qn('w:tblInd'))
    ti.set(docx_qn('w:w'), str(TABLE_INDENT))
    ti.set(docx_qn('w:type'), 'dxa')

    borders = lxml_etree.SubElement(tbl_pr, docx_qn('w:tblBorders'))
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = lxml_etree.SubElement(borders, docx_qn(f'w:{border_name}'))
        el.set(docx_qn('w:val'), 'none')
        el.set(docx_qn('w:sz'), '0')
        el.set(docx_qn('w:space'), '0')
        el.set(docx_qn('w:color'), 'auto')

    top_border = borders.find(docx_qn('w:top'))
    top_border.set(docx_qn('w:val'), 'single')
    top_border.set(docx_qn('w:sz'), '8')

    bottom_border = borders.find(docx_qn('w:bottom'))
    bottom_border.set(docx_qn('w:val'), 'single')
    bottom_border.set(docx_qn('w:sz'), '8')

    tl = lxml_etree.SubElement(tbl_pr, docx_qn('w:tblLayout'))
    tl.set(docx_qn('w:type'), 'fixed')
    tcm = lxml_etree.SubElement(tbl_pr, docx_qn('w:tblCellMar'))
    for side in ['left', 'right']:
        s = lxml_etree.SubElement(tcm, docx_qn(f'w:{side}'))
        s.set(docx_qn('w:w'), '0')
        s.set(docx_qn('w:type'), 'dxa')

    for ri, row_data in enumerate(table_data):
        col_idx = 0
        for cell_data in row_data:
            if col_idx >= ncols:
                break
            if cell_data.get('vmerge_continue'):
                col_idx += cell_data.get('gridspan', 1)
                continue
            cell = table.cell(ri, col_idx)
            for p in cell.paragraphs:
                p.clear()
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = para.paragraph_format
            pf.line_spacing = 1.0

            if ri == 0:
                tc_pr = cell._tc.find(docx_qn('w:tcPr'))
                if tc_pr is None:
                    tc_pr = lxml_etree.SubElement(cell._tc, docx_qn('w:tcPr'))
                tc_borders = lxml_etree.SubElement(tc_pr, docx_qn('w:tcBorders'))
                bb = lxml_etree.SubElement(tc_borders, docx_qn('w:bottom'))
                bb.set(docx_qn('w:val'), 'single')
                bb.set(docx_qn('w:sz'), '4')
                bb.set(docx_qn('w:space'), '0')
                bb.set(docx_qn('w:color'), 'auto')

            if cell_data.get('runs'):
                for rd in cell_data['runs']:
                    if not rd['text']:
                        continue
                    run = para.add_run(rd['text'])
                    b = rd['bold'] or (ri == 0)
                    _set_run_font(run, size=DEFAULT_SIZE, bold=b,
                                  italic=rd['italic'],
                                  superscript=rd.get('superscript', False),
                                  subscript=rd.get('subscript', False))
            else:
                run = para.add_run(cell_data.get('text', '').strip())
                _set_run_font(run, size=DEFAULT_SIZE, bold=(ri == 0))

            gs = cell_data.get('gridspan', 1)
            if gs > 1 and col_idx + gs <= ncols:
                cell.merge(table.cell(ri, col_idx + gs - 1))
            col_idx += gs

    return table


def build(items, output_path, ris_data=None, zotero_enabled=False,
          use_crossref=False, progress_callback=None):
    doc = DocxDocument()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.50)
    section.bottom_margin = Cm(1.60)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    # Default style
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = DEFAULT_SIZE
    style.font.color.rgb = FONT_COLOR

    # Register named MDPI styles
    register_mdpi_styles(doc)

    # Article type
    para = doc.add_paragraph()
    _apply_direct_style(para, space_before=Pt(12),
                        alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run('Article')
    _set_run_font(run, size=DEFAULT_SIZE, italic=True)

    # Find title
    title_elem = None
    for elem in items:
        if elem['type'] in ('abstract_heading', 'heading1'):
            break
        if elem['type'] == 'paragraph':
            if not title_elem:
                title_elem = elem

    # Title
    title_text = title_elem['text'] if title_elem else '[Title]'
    para = doc.add_paragraph()
    _apply_direct_style(para, space_after=Pt(12), line_spacing_pt=Pt(12))
    run = para.add_run(title_text)
    _set_run_font(run, size=Pt(18), bold=True)

    # Author placeholder
    _build_author_placeholder(doc)

    # Resolve references through the engine
    ref_items = [item for item in items if item['type'] == 'reference']
    resolved_refs = resolve_references(
        ref_items, ris_data=ris_data, use_crossref=use_crossref,
        progress_callback=progress_callback)
    ref_index = 0

    # Main content
    after_heading = False
    for item in items:
        itype = item['type']

        if item is title_elem:
            continue
        if itype == 'abstract_heading':
            continue

        if itype == 'abstract_text':
            _add_runs_para(doc, 'Normal', item['runs'], bold_prefix='Abstract: ',
                           font_size=DEFAULT_SIZE)
            # Apply abstract formatting directly
            para = doc.paragraphs[-1]
            _apply_direct_style(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                space_before=Pt(12), space_after=Pt(6),
                                line_spacing_pt=Pt(14), left_indent=LEFT_INDENT)
            after_heading = False
            continue

        if itype == 'keywords':
            kw_text = item['text'].replace('Keywords:', '').replace(
                'Keywords :', '').strip()
            para = doc.add_paragraph()
            _apply_direct_style(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                space_before=Pt(12), line_spacing_pt=Pt(14),
                                left_indent=LEFT_INDENT)
            run = para.add_run('Keywords: ')
            _set_run_font(run, size=DEFAULT_SIZE, bold=True)
            run = para.add_run(kw_text)
            _set_run_font(run, size=DEFAULT_SIZE)
            after_heading = False
            continue

        if itype == 'heading1':
            para = doc.add_paragraph()
            _apply_direct_style(para, space_before=Pt(12), space_after=Pt(3),
                                line_spacing_pt=Pt(14), left_indent=LEFT_INDENT,
                                outline_level=0)
            run = para.add_run(item['text'])
            _set_run_font(run, size=Pt(12), bold=True)
            after_heading = True
            continue

        if itype == 'heading2':
            para = doc.add_paragraph()
            _apply_direct_style(para, space_before=Pt(3), space_after=Pt(3),
                                line_spacing_pt=Pt(14), left_indent=LEFT_INDENT,
                                outline_level=1)
            run = para.add_run(item['text'])
            _set_run_font(run, size=DEFAULT_SIZE, italic=True)
            after_heading = True
            continue

        if itype == 'heading3':
            para = doc.add_paragraph()
            _apply_direct_style(para, space_before=Pt(3), space_after=Pt(3),
                                line_spacing_pt=Pt(14), left_indent=LEFT_INDENT,
                                outline_level=2)
            run = para.add_run(item['text'])
            _set_run_font(run, size=DEFAULT_SIZE)
            after_heading = True
            continue

        if itype == 'references_heading':
            para = doc.add_paragraph()
            _apply_direct_style(para, space_before=Pt(12), space_after=Pt(3),
                                line_spacing_pt=Pt(14), left_indent=LEFT_INDENT,
                                outline_level=0)
            run = para.add_run(item['text'])
            _set_run_font(run, size=Pt(12), bold=True)
            after_heading = True
            continue

        if itype == 'paragraph':
            para = doc.add_paragraph()
            fi = Pt(0) if after_heading else Pt(21.25)
            _apply_direct_style(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                line_spacing_pt=Pt(14), left_indent=LEFT_INDENT,
                                first_indent=fi)
            for rd in item['runs']:
                if not rd['text']:
                    continue
                run = para.add_run(rd['text'])
                _set_run_font(run, size=DEFAULT_SIZE, bold=rd.get('bold'),
                              italic=rd.get('italic'),
                              superscript=rd.get('superscript', False),
                              subscript=rd.get('subscript', False))
            after_heading = False
            continue

        if itype == 'equation':
            para = doc.add_paragraph()
            _apply_direct_style(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                line_spacing_pt=Pt(14), left_indent=LEFT_INDENT,
                                first_indent=Pt(0))
            for rd in item['runs']:
                if not rd['text']:
                    continue
                run = para.add_run(rd['text'])
                _set_run_font(run, size=DEFAULT_SIZE, bold=rd.get('bold'),
                              italic=rd.get('italic'))
            after_heading = False
            continue

        if itype == 'table_caption':
            text = item['text']
            match = re.match(r'^Table\s+(\d+)\.?\s*(.*)', text)
            if match:
                num = match.group(1)
                desc = match.group(2)
                para = doc.add_paragraph()
                try:
                    para.style = doc.styles['MDPI41tablecaption']
                except KeyError:
                    pass
                add_caption_with_seq(para, 'Table', num, description=desc,
                                     font_name=FONT_NAME, font_size=Pt(9),
                                     font_color=FONT_COLOR, bold_label=True)
            else:
                _add_styled_para(doc, 'MDPI41tablecaption', text, font_size=Pt(9))
            after_heading = False
            continue

        if itype == 'table':
            _add_three_line_table(doc, item['rows'])
            after_heading = False
            continue

        if itype == 'table_footer':
            _add_styled_para(doc, 'MDPI43tablefooter', item['text'], font_size=Pt(9))
            after_heading = False
            continue

        if itype == 'figure_placeholder':
            text = item['text']
            match = re.match(r'^\[?\s*Figure\s+(\d+)\]?\.?\s*(.*)', text, re.IGNORECASE)
            if match:
                num = match.group(1)
                desc = match.group(2)
                para = doc.add_paragraph()
                try:
                    para.style = doc.styles['MDPI51figurecaption']
                except KeyError:
                    pass
                add_caption_with_seq(para, 'Figure', num, description=desc,
                                     font_name=FONT_NAME, font_size=Pt(9),
                                     font_color=FONT_COLOR, bold_label=True)
            else:
                _add_runs_para(doc, 'MDPI51figurecaption', item['runs'], font_size=Pt(9))
            after_heading = False
            continue

        if itype == 'reference':
            if ref_index < len(resolved_refs):
                rr = resolved_refs[ref_index]
                ref_index += 1

                para = doc.add_paragraph()
                try:
                    para.style = doc.styles['MDPIBibliography']
                except KeyError:
                    pass

                if zotero_enabled and rr.metadata:
                    csl_json = ris_to_csl_json(rr.metadata, rr.index)
                    visible = ''.join(r['text'] for r in rr.formatted_runs)
                    add_zotero_citation_field(para, csl_json, visible,
                                              font_name=FONT_NAME,
                                              font_size=DEFAULT_SIZE)
                else:
                    for rd in rr.formatted_runs:
                        run = para.add_run(rd['text'])
                        _set_run_font(run, size=DEFAULT_SIZE,
                                      bold=rd.get('bold'),
                                      italic=rd.get('italic'),
                                      superscript=rd.get('superscript', False),
                                      subscript=rd.get('subscript', False))
            after_heading = False
            continue

    doc.save(output_path)
    return output_path
