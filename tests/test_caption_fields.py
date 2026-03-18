import pytest
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn as docx_qn


def test_add_caption_with_seq():
    from caption_fields import add_caption_with_seq
    doc = Document()
    para = doc.add_paragraph()
    bm_name = add_caption_with_seq(para, 'Table', 1, description='Sample data',
                                    font_size=Pt(9))
    # Returns a bookmark name
    assert bm_name.startswith('_RefTable1_')
    # Check that the paragraph has fldChar elements (SEQ field)
    p_xml = para._p
    fld_chars = p_xml.findall(f'.//{docx_qn("w:fldChar")}')
    assert len(fld_chars) == 3  # begin, separate, end
    # Check instrText contains SEQ
    instr = p_xml.findall(f'.//{docx_qn("w:instrText")}')
    assert any('SEQ Table' in (el.text or '') for el in instr)
    # Check visible text includes number
    texts = [t.text for t in p_xml.iter(docx_qn('w:t')) if t.text]
    full = ''.join(texts)
    assert 'Table' in full
    assert '1' in full
    assert 'Sample data' in full
    # Check bookmarkStart and bookmarkEnd present
    bm_starts = p_xml.findall(f'.//{docx_qn("w:bookmarkStart")}')
    bm_ends = p_xml.findall(f'.//{docx_qn("w:bookmarkEnd")}')
    assert len(bm_starts) == 1
    assert len(bm_ends) == 1
    assert bm_starts[0].get(docx_qn('w:name')) == bm_name
    # Check w:noProof on the number run
    no_proofs = p_xml.findall(f'.//{docx_qn("w:noProof")}')
    assert len(no_proofs) >= 1


def test_add_figure_caption_with_seq():
    from caption_fields import add_caption_with_seq
    doc = Document()
    para = doc.add_paragraph()
    bm_name = add_caption_with_seq(para, 'Figure', 3, description='A nice plot',
                                    font_size=Pt(10))
    assert bm_name.startswith('_RefFigure3_')
    p_xml = para._p
    instr = p_xml.findall(f'.//{docx_qn("w:instrText")}')
    assert any('SEQ Figure' in (el.text or '') for el in instr)


def test_ris_to_csl_json():
    from caption_fields import ris_to_csl_json
    import json
    ref = {
        'type': 'JOUR',
        'authors': ['Shannon, Claude E.'],
        'title': 'A Mathematical Theory of Communication',
        'journal': 'Bell System Technical Journal',
        'year': '1948',
        'volume': '27',
        'start_page': '379',
        'end_page': '423',
        'doi': '10.1002/j.1538-7305.1948.tb01338.x',
        'issue': '', 'publisher': '', 'place': '', 'edition': '',
        'editors': [], 'url': '', 'keywords': [], 'isbn': '', 'abstract': '',
    }
    result = ris_to_csl_json(ref, 1)
    data = json.loads(result)
    assert 'citationID' in data
    assert data['citationItems'][0]['itemData']['title'] == 'A Mathematical Theory of Communication'
    assert data['citationItems'][0]['itemData']['type'] == 'article-journal'


def test_add_zotero_citation_field():
    from caption_fields import add_zotero_citation_field
    doc = Document()
    para = doc.add_paragraph()
    add_zotero_citation_field(para, '{"test": true}', '[1]',
                               font_size=Pt(10))
    p_xml = para._p
    fld_chars = p_xml.findall(f'.//{docx_qn("w:fldChar")}')
    assert len(fld_chars) == 3
    instr = p_xml.findall(f'.//{docx_qn("w:instrText")}')
    assert any('ZOTERO_ITEM' in (el.text or '') for el in instr)
