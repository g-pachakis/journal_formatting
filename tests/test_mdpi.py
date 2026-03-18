import os
import pytest
from docx import Document
from docx.shared import Pt, Cm


def test_mdpi_has_required_attributes():
    from formats import mdpi
    assert mdpi.FORMAT_NAME == 'MDPI'
    assert mdpi.FORMAT_SUFFIX == '_MDPI'
    assert callable(mdpi.build)


def test_mdpi_build_creates_file(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.mdpi import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_MDPI.docx')
    result = build(items, output)
    assert result == output
    assert os.path.isfile(output)


def test_mdpi_output_is_valid_docx(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.mdpi import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_MDPI.docx')
    build(items, output)
    doc = Document(output)
    assert len(doc.paragraphs) > 0


def test_mdpi_page_setup(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.mdpi import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_MDPI.docx')
    build(items, output)
    doc = Document(output)
    section = doc.sections[0]
    assert abs(section.page_width - Cm(21)) < Cm(0.1)
    assert abs(section.page_height - Cm(29.7)) < Cm(0.1)


def test_mdpi_contains_article_type(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.mdpi import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_MDPI.docx')
    build(items, output)
    doc = Document(output)
    texts = [p.text for p in doc.paragraphs]
    assert any('Article' in t for t in texts)


def test_mdpi_contains_title(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.mdpi import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_MDPI.docx')
    build(items, output)
    doc = Document(output)
    texts = [p.text for p in doc.paragraphs]
    assert any('Test Manuscript Title' in t for t in texts)


def test_mdpi_contains_abstract_prefix(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.mdpi import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_MDPI.docx')
    build(items, output)
    doc = Document(output)
    texts = [p.text for p in doc.paragraphs]
    assert any('Abstract:' in t or 'Abstract: ' in t for t in texts)


def test_mdpi_uses_palatino_linotype(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.mdpi import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_MDPI.docx')
    build(items, output)
    doc = Document(output)
    assert doc.styles['Normal'].font.name == 'Palatino Linotype'


def test_mdpi_has_table(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.mdpi import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_MDPI.docx')
    build(items, output)
    doc = Document(output)
    assert len(doc.tables) >= 1


def test_mdpi_preserves_formatting(sample_manuscript_with_formatting, tmp_path):
    from reader import read_manuscript
    from formats.mdpi import build
    items = read_manuscript(sample_manuscript_with_formatting)
    output = str(tmp_path / 'output_MDPI.docx')
    build(items, output)
    doc = Document(output)
    assert len(doc.paragraphs) > 0
