"""End-to-end integration test: reader -> plugin -> valid output."""

import os
import pytest
from docx import Document


@pytest.fixture(params=['MDPI', 'Elsevier'])
def format_name(request):
    return request.param


def test_full_pipeline(sample_manuscript, format_name, tmp_path):
    from reader import read_manuscript
    from formats import get_formats

    formats = get_formats()
    assert format_name in formats

    plugin = formats[format_name]
    items = read_manuscript(sample_manuscript)
    assert len(items) > 0

    output = str(tmp_path / f'output{plugin.FORMAT_SUFFIX}.docx')
    result = plugin.build(items, output)

    assert result == output
    assert os.path.isfile(output)

    doc = Document(output)
    assert len(doc.paragraphs) > 0

    all_text = ' '.join(p.text for p in doc.paragraphs)
    assert 'Introduction' in all_text or '1.' in all_text


def test_full_pipeline_with_formatting(sample_manuscript_with_formatting,
                                        format_name, tmp_path):
    from reader import read_manuscript
    from formats import get_formats

    plugin = get_formats()[format_name]
    items = read_manuscript(sample_manuscript_with_formatting)
    output = str(tmp_path / f'formatted{plugin.FORMAT_SUFFIX}.docx')
    plugin.build(items, output)

    doc = Document(output)
    assert len(doc.paragraphs) > 0
