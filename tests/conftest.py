import os
import pytest
from docx import Document
from docx.shared import Pt


@pytest.fixture
def sample_manuscript(tmp_path):
    """Create a minimal manuscript .docx for testing."""
    doc = Document()

    # Title paragraph (before any heading)
    p = doc.add_paragraph('Test Manuscript Title')

    # Abstract heading
    doc.add_heading('Abstract', level=1)

    # Abstract text
    doc.add_paragraph('This is the abstract text of the manuscript.')

    # Keywords
    doc.add_paragraph('Keywords: keyword1; keyword2; keyword3')

    # Section heading
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('This is the introduction paragraph.')
    doc.add_paragraph('This is the second paragraph of the introduction.')

    # Subsection
    doc.add_heading('1.1. Background', level=2)
    doc.add_paragraph('Background text here.')

    # Sub-subsection
    doc.add_heading('1.1.1. Details', level=3)
    doc.add_paragraph('Detail text here.')

    # Table caption
    doc.add_paragraph('Table 1. Sample data table')

    # Table
    table = doc.add_table(rows=3, cols=3)
    headers = ['Column A', 'Column B', 'Column C']
    for i, h in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
    for r in range(1, 3):
        for c in range(3):
            table.rows[r].cells[c].text = f'Data {r},{c}'

    # Table footer
    doc.add_paragraph('* Table footnote explaining something.')

    # Figure placeholder
    doc.add_paragraph('Figure 1. A sample figure caption')

    # Equation
    doc.add_paragraph('The relationship is given by y = mx + b (Eq. 1)')

    # References heading
    doc.add_heading('References', level=1)
    doc.add_paragraph('[1] Author A. Title of paper. Journal, 2024.')
    doc.add_paragraph('[2] Author B. Another paper. Journal, 2023.')

    path = tmp_path / 'test_manuscript.docx'
    doc.save(str(path))
    return str(path)


@pytest.fixture
def sample_manuscript_with_formatting(tmp_path):
    """Create a manuscript with bold/italic/super/subscript runs."""
    doc = Document()

    doc.add_heading('Abstract', level=1)

    p = doc.add_paragraph()
    p.add_run('Normal text, ')
    bold_run = p.add_run('bold text')
    bold_run.bold = True
    p.add_run(', ')
    italic_run = p.add_run('italic text')
    italic_run.italic = True
    p.add_run(', H')
    sub_run = p.add_run('2')
    sub_run.font.subscript = True
    p.add_run('O and x')
    sup_run = p.add_run('2')
    sup_run.font.superscript = True

    doc.add_paragraph('Keywords: water; energy')

    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('Body paragraph.')

    path = tmp_path / 'formatted_manuscript.docx'
    doc.save(str(path))
    return str(path)
