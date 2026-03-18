import os
import pytest
from docx import Document


def test_elsevier_has_required_attributes():
    from formats import elsevier
    assert elsevier.FORMAT_NAME == 'Elsevier'
    assert elsevier.FORMAT_SUFFIX == '_Elsevier'
    assert callable(elsevier.build)


def test_elsevier_build_creates_file(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.elsevier import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_Elsevier.docx')
    result = build(items, output)
    assert result == output
    assert os.path.isfile(output)


def test_elsevier_output_is_valid_docx(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.elsevier import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_Elsevier.docx')
    build(items, output)
    doc = Document(output)
    assert len(doc.paragraphs) > 0


def test_elsevier_page_setup(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.elsevier import build
    from docx.shared import Cm
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_Elsevier.docx')
    build(items, output)
    doc = Document(output)
    section = doc.sections[0]
    assert abs(section.page_width - Cm(21)) < Cm(0.1)
    assert abs(section.page_height - Cm(29.7)) < Cm(0.1)
    assert abs(section.left_margin - Cm(2.5)) < Cm(0.1)


def test_elsevier_contains_title(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.elsevier import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_Elsevier.docx')
    build(items, output)
    doc = Document(output)
    texts = [p.text for p in doc.paragraphs]
    assert any('Test Manuscript Title' in t for t in texts)


def test_elsevier_contains_highlights_section(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.elsevier import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_Elsevier.docx')
    build(items, output)
    doc = Document(output)
    texts = [p.text for p in doc.paragraphs]
    assert any('Highlights' in t for t in texts)


def test_elsevier_contains_references(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.elsevier import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_Elsevier.docx')
    build(items, output)
    doc = Document(output)
    texts = [p.text for p in doc.paragraphs]
    assert any('References' in t for t in texts)


def test_elsevier_has_table(sample_manuscript, tmp_path):
    from reader import read_manuscript
    from formats.elsevier import build
    items = read_manuscript(sample_manuscript)
    output = str(tmp_path / 'output_Elsevier.docx')
    build(items, output)
    doc = Document(output)
    assert len(doc.tables) >= 2
